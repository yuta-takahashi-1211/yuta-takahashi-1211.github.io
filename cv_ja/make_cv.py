from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

doc = Document()

# ページ余白設定
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# デフォルトフォント設定
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '游明朝')

def set_font(run, size=10.5, bold=False, name='Times New Roman', east='游明朝'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), east)

def add_heading(doc, text, size=12, bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, east='游ゴシック')
    return p

def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10.5, bold=True, east='游ゴシック')
    return p

def add_body(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run)
    return p

# ─────────────────────────────────────────
# 氏名・肩書き・所属
# ─────────────────────────────────────────
add_heading(doc, '【氏名】')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('高橋 悠太（たかはし ゆうた / Yuta Takahashi）')
set_font(run)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('准教授')
set_font(run)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('研究領域：マクロ経済学')
set_font(run)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run('所属部門：実証経済部門')
set_font(run)

# ─────────────────────────────────────────
# 学歴・職歴
# ─────────────────────────────────────────
add_heading(doc, '【学歴・職歴】')

education = [
    ('2009年 3月', '京都大学 経済学部 卒業', '学士（経済学）'),
    ('2011年 3月', '京都大学大学院 経済学研究科 修了', '修士（経済学）'),
    ('2018年 8月', 'Northwestern University 卒業', 'Ph.D. in Economics'),
]

career = [
    ('2018年 9月 〜 2020年 3月', '一橋大学 経済研究所', '非常勤研究員（ポスドク）'),
    ('2020年 4月 〜 2025年 8月', '一橋大学 経済研究所', '講師'),
    ('2025年 9月 〜 現在', '大阪大学 社会経済研究所', '准教授'),
]

for year, inst, degree in education:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'{year}　{inst}　{degree}')
    set_font(run)

doc.add_paragraph()

for period, inst, pos in career:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'{period}　{inst}　{pos}')
    set_font(run)

# ─────────────────────────────────────────
# 主な発表論文名・著書名等
# ─────────────────────────────────────────
add_heading(doc, '【主な発表論文名・著書名等】')

add_subheading(doc, '発表論文（査読付き）')

papers = [
    (
        '"Universal Gravity,"',
        'with Treb Allen and Costas Arkolakis,',
        'Journal of Political Economy, 128(2): 393–433, 2020.'
    ),
    (
        '"A Parsimonious Model for Zero Inflation at the Zero Lower Bound,"',
        'with Naoki Takayama,',
        'Economics Letters, Volume 247, February 2025.'
    ),
    (
        '"An Experiment on a Multi-Period Beauty Contest Game,"',
        'with Nobuyuki Hanaki,',
        'Forthcoming at Experimental Economics, 2025.'
    ),
]

for i, (title, coauthor, journal) in enumerate(papers, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r1 = p.add_run(f'[{i}]  ')
    set_font(r1)
    r2 = p.add_run(title)
    set_font(r2, bold=True)
    r3 = p.add_run(f'\n      {coauthor}\n      {journal}')
    set_font(r3)

add_subheading(doc, '未刊行論文')

wps = [
    (
        '"Does Expected Inflation Matter? Evidence from Japanese Value-Added Tax Hikes"',
        'with Naoki Takayama, 2024.'
    ),
    (
        '"Global Technology Stagnation,"',
        'with Naoki Takayama, 2024.'
    ),
    (
        '"Hidden Stagflation,"',
        'with Naoki Takayama, 2024.'
    ),
    (
        '"Anchoring Inflation Expectation,"',
        'with Lawrence Christiano, 2020.'
    ),
]

for i, (title, info) in enumerate(wps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    r1 = p.add_run(f'[{i}]  ')
    set_font(r1)
    r2 = p.add_run(title)
    set_font(r2, bold=True)
    r3 = p.add_run(f'\n      {info}')
    set_font(r3)

# ─────────────────────────────────────────
# その他
# ─────────────────────────────────────────
add_heading(doc, '【その他】')

add_subheading(doc, '(1) 学会・セミナー・講演')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Seminar in Macroeconomics 共同オーガナイザー（一橋大学 経済研究所・京都大学 経済研究所 共同運営セミナー）')
set_font(run)

add_subheading(doc, '(2) レフェリー経験')
referee_text = (
    'Canadian Journal of Economics, '
    'Economic Journal, '
    'Economic Modelling, '
    'Economic Theory, '
    'Economics Letters, '
    'European Economic Review, '
    'Games and Economic Behavior, '
    'Hitotsubashi Journal of Economics, '
    'International Economic Review, '
    'Journal of Applied Economics, '
    'Journal of Economic Behavior & Organization, '
    'Journal of International Economics, '
    'Journal of Macroeconomics, '
    'Journal of Political Economy, '
    'Journal of the Japanese and International Economies, '
    'Quantitative Economics, '
    'Review of Economic Dynamics, '
    'Review of Economic Studies, '
    '経済研究'
)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run(referee_text)
set_font(run)

# ─────────────────────────────────────────
# 講演・口頭発表等
# ─────────────────────────────────────────
add_heading(doc, '【講演・口頭発表等】')

presentations = {
    2025: [
        ('5月', 'The 2nd International SPACE Workshop', '招待'),
        ('5月', '財務総合政策研究所セミナー', '招待'),
        ('8月', 'SWET 2025 (Macroeconomics, Financial Economics Session)', ''),
        ('8月', '2025 World Congress of the Econometric Society', ''),
        ('9月', 'The 1st Workshop of the East and South-East Asian Macroeconomic Society', ''),
        ('10月', 'International Workshop on Experiments in Macroeconomics and Finance', '招待'),
        ('11月', 'WAMS 2025', ''),
        ('12月', 'OzMac 2025', ''),
    ],
    2024: [
        ('3月', 'HSI2023-The 9th Hitotsubashi Summer Institute', ''),
        ('3月', 'ANU-AJRC, Hitotsubashi, Keio and University of Tokyo Conference on Frontiers in Macroeconomics', ''),
    ],
    2023: [
        ('7月', '2023 Asian Meeting of the Econometric Society in East and Southeast Asia (AMES)', ''),
        ('9月', '日本経済学会 2023 年度秋季大会', '招待'),
        ('12月', 'European Winter Meeting of the Econometric Society 2023', ''),
    ],
    2021: [
        ('1月', 'マクロ経済学ワークショップ（一橋大学）', ''),
        ('4月', 'マクロ経済学ワークショップ 2021（東京大学 CIRJE）', ''),
        ('4月', 'マクロ金融ワークショップ（一橋大学）', ''),
        ('5月', '第1回マクロ税制セミナー（学習院大学）', ''),
        ('6月', '早稲田大学セミナー', ''),
        ('6月', 'マクロ経済学ワークショップ（慶應大学）', ''),
        ('7月', 'SED', ''),
        ('7月', '六甲フォーラム（神戸大学）', ''),
        ('9月', 'GSE-OSIPP-ISER Joint Conference in Economics（大阪大学）', ''),
        ('12月', '7th Annual CIGS End of Year Macroeconomics Conference', ''),
    ],
    2020: [
        ('1月', '東北大学現代経済学研究会', ''),
        ('9月', '経済成長と構造変化に関する長期分析研究会', ''),
        ('10月', 'ISER Virtual Workshop in Macro/International Economics', ''),
    ],
    2019: [
        ('5月', 'CIGS Conference on Macroeconomic Theory and Policy 2019', ''),
        ('9月', 'Monetary Policy Workshop', ''),
    ],
}

for year in sorted(presentations.keys(), reverse=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f'{year}年')
    set_font(run, bold=True)

    for month, venue, invite in presentations[year]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        text = f'{month}　{venue}'
        if invite:
            text += f'　（{invite}）'
        run = p.add_run(text)
        set_font(run)

# 保存
out_path = 'cv_ja/cv_ja.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
